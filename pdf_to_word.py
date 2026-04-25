import streamlit as st
import pdfplumber
import re
import os
import shutil
import io
import base64
import pandas as pd
import json
import csv
import time
import concurrent.futures
import fitz
from PIL import Image
import pytesseract
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage
from google.genai.types import HarmBlockThreshold, HarmCategory
from dotenv import load_dotenv

load_dotenv()

st.set_page_config(page_title="AI PDF to Word", page_icon="📄", layout="wide")
st.title("📄 Trợ lý AI: Biên tập Đề Thi & Tài Liệu")


GEMINI_MODEL_NAME = "gemini-2.5-flash"

# Lưu ý: API safety_settings hiện chỉ chấp nhận các category "core" bên dưới.
GEMINI_SAFETY_SETTINGS = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_CIVIC_INTEGRITY: HarmBlockThreshold.BLOCK_NONE,
}


def build_gemini_model(temperature=0.1):
    return ChatGoogleGenerativeAI(
        model=GEMINI_MODEL_NAME,
        temperature=temperature,
        safety_settings=GEMINI_SAFETY_SETTINGS,
        max_output_tokens=8192,
    )


def clean_ai_output(text):
    """Làm sạch output từ AI: loại bỏ backslash escape thừa."""
    if not text:
        return text
    # Loại bỏ backslash trước dấu gạch dưới: \_ → _
    return re.sub(r"\\_", "_", text)


def _configure_tesseract_cmd_if_possible():
    env_cmd = os.getenv("TESSERACT_CMD")
    candidates = [
        env_cmd,
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
    ]
    for cmd in candidates:
        if cmd and os.path.exists(cmd):
            pytesseract.pytesseract.tesseract_cmd = cmd
            return


def is_tesseract_available():
    try:
        _configure_tesseract_cmd_if_possible()
        if shutil.which(
            pytesseract.pytesseract.tesseract_cmd
        ) is None and not os.path.exists(pytesseract.pytesseract.tesseract_cmd):
            return False
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def ocr_images_with_tesseract(images):
    if not is_tesseract_available():
        raise RuntimeError("TESSERACT_NOT_AVAILABLE")

    texts = []
    for img in images:
        # Ưu tiên OCR tiếng Việt + tiếng Anh; nếu thiếu data thì rơi về tiếng Anh.
        try:
            texts.append(pytesseract.image_to_string(img, lang="vie+eng"))
        except Exception:
            texts.append(pytesseract.image_to_string(img, lang="eng"))
    return "\n\n".join([t for t in texts if t])


# ==========================================
# CÁC HÀM XỬ LÝ PDF
# ==========================================
def extract_text_from_upload(uploaded_file):
    text = ""
    pdf_bytes = b""
    try:
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()
        pdf_buffer = io.BytesIO(pdf_bytes)
        with pdfplumber.open(pdf_buffer) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
    except Exception as e:
        st.error(f"❌ Lỗi đọc chữ PDF (pdfplumber): {e}")

    # Fallback: nhiều PDF pdfplumber trích xuất rỗng nhưng PyMuPDF đọc được text.
    if not text.strip() and pdf_bytes:
        try:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page in doc:
                extracted = page.get_text("text")
                if extracted:
                    text += extracted + "\n"
        except Exception as e:
            st.error(f"❌ Lỗi đọc chữ PDF (PyMuPDF): {e}")
    return text


def extract_images_from_upload(uploaded_file):
    images = []
    pdf_bytes = b""
    try:
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()

        # Ưu tiên PyMuPDF để render ổn định (tránh nền đen do alpha/transparent)
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        zoom = 150 / 72  # 150 dpi
        matrix = fitz.Matrix(zoom, zoom)
        for page in doc:
            pix = page.get_pixmap(matrix=matrix, alpha=True)
            img_rgba = Image.frombytes("RGBA", (pix.width, pix.height), pix.samples)
            white_bg = Image.new("RGB", img_rgba.size, (255, 255, 255))
            white_bg.paste(img_rgba, mask=img_rgba.split()[3])
            images.append(white_bg)
    except Exception as e:
        # Fallback cuối: dùng pdfplumber render ảnh
        try:
            uploaded_file.seek(0)
            pdf_buffer = io.BytesIO(pdf_bytes or uploaded_file.read())
            with pdfplumber.open(pdf_buffer) as pdf:
                for page in pdf.pages:
                    img = page.to_image(resolution=150).original
                    images.append(img)
        except Exception as e2:
            st.error(f"❌ Lỗi chuyển PDF thành ảnh: {e2}")
    return images


# ==========================================
# HÀM AI - CẬP NHẬT LUẬT MỀM DẺO CHO BẢNG
# ==========================================
EDUCATIONAL_SYSTEM_PROMPT = """Bạn là một Chuyên gia Biên tập Đề thi và Tài liệu Giáo dục Tiếng Anh.
Nhiệm vụ của bạn là định dạng lại văn bản cực kỳ CHÍNH XÁC theo cấu trúc chuẩn.
LƯU Ý TỐI QUAN TRỌNG:
1. MARKDOWN: Dùng dấu sao đôi để in đậm (**chữ**). KHÔNG dùng thẻ HTML.
2. BẢNG TỪ VỰNG (Vocabulary): Giữ nguyên đầy đủ các cột (STT, Word, Part of speech, Pronunciation, Meaning). Định dạng bằng markdown table.
3. BẢNG CẤU TRÚC (Structures): Giữ nguyên đầy đủ các cột. Định dạng bằng markdown table.
4. QUIZ (Điền từ): Giữ nguyên câu hỏi và các lựa chọn (dạng A/B hoặc A/B/C).
5. GRAMMAR: Giữ nguyên nội dung ngữ pháp, ví dụ, công thức.
6. CÂU HỎI TRẮC NGHIỆM (Multiple Choice):
   - Bắt buộc in đậm chữ Question và số (VD: **Question 1.**). KHÔNG in đậm nội dung câu hỏi.
   - Bắt buộc in đậm các chữ cái đáp án (VD: **A.**, **B.**, **C.**, **D.**).
   - Các đáp án phải nằm trên cùng một dòng.
7. ĐỌC HIỂU (Reading Comprehension):
   - Giữ nguyên sự liền mạch của đoạn văn, không tự ý ngắt dòng giữa câu.
   - Chỗ điền từ (blanks) dùng dấu gạch dưới: _____ (5 dấu).
8. LỌC RÁC: Xóa các thông tin như Link web, số điện thoại, tên giáo viên ở đầu/cuối tài liệu.
9. KHÔNG bỏ sót bất kỳ phần nào: Vocabulary, Structures, Quiz, Grammar, Practice, Reading.
10. KHÔNG giao tiếp, KHÔNG giải thích. CHỈ TRẢ VỀ văn bản đã được biên tập."""


def process_text_with_ai(raw_text, user_prompt):
    model = build_gemini_model(temperature=0.1)
    full_prompt = f"{EDUCATIONAL_SYSTEM_PROMPT}\n\nYÊU CẦU CỦA NGƯỜI DÙNG:\n{user_prompt}\n\nNỘI DUNG TÀI LIỆU GỐC:\n<document>\n{raw_text}\n</document>"
    return clean_ai_output(model.invoke(full_prompt).content)


# --- HÀM XỬ LÝ SINGLE PAGE ĐÃ ĐƯỢC VIẾT LẠI TRONG process_vision_with_ai ---
# XÓA HÀM CŨ ĐỂ TRÁNH RUNTIME ERROR


# --- Hàm chính xử lý song song ---
def process_vision_with_ai(images, user_prompt):
    total_pages = len(images)

    results_map = {}
    debug_map = {}

    # XỬ LÝ MỖI TRANG RIÊNG BIỆT - MỖI TRANG 1 REQUEST (FIX BỎ SÓT TRANG)
    for page_num in range(total_pages):
        try:
            print(f"Processing page {page_num + 1}/{total_pages}...")

            # Gửi MỖI TRANG RIÊNG LẺ, KHÔNG GỬI NHIỀU ẢNH CÙNG LÚC
            current_image = images[page_num]

            full_result = ""

            # Retry tối đa 3 lần nếu output ngắn
            for attempt in range(3):
                # TẠO MODEL MỚI MỖI LẦN REQUEST
                model = build_gemini_model(temperature=0.1)

                # Encode image to base64
                img_buffer = io.BytesIO()
                current_image.save(img_buffer, format="JPEG")
                img_base64 = base64.b64encode(img_buffer.getvalue()).decode("utf-8")

                content = [
                    {
                        "type": "text",
                        "text": f"{EDUCATIONAL_SYSTEM_PROMPT}\n\nYêu cầu: {user_prompt}\n\nRẤT QUAN TRỌNG: ĐÂY LÀ TRANG {page_num + 1} / {total_pages}. HÃY TRÍCH XUẤT TOÀN BỘ NỘI DUNG CỦA TRANG NÀY ĐƠN ĐỘC, KHÔNG BỎ SÓT BẤT KỲ NỘI DUNG NÀO.",
                    },
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{img_base64}"},
                    },
                ]

                try:
                    resp = model.invoke([HumanMessage(content=content)])
                    text = clean_ai_output(str(getattr(resp, "content", "") or ""))
                    text_len = len(text.strip())

                    print(f"  Attempt {attempt + 1}: {text_len} chars")

                    if text_len >= 300:
                        full_result = text
                        break  # Đủ nội dung rồi

                    time.sleep(3 + attempt * 2)

                except Exception as e:
                    print(f"  Error: {str(e)[:100]}")
                    time.sleep(5)
                    continue

            results_map[page_num] = full_result
            debug_map[page_num] = {
                "page": page_num + 1,
                "text_len": len(full_result),
                "attempts": attempt + 1,
            }

            # Delay 4 giây giữa các trang
            time.sleep(4)

        except Exception as e:
            results_map[page_num] = ""
            debug_map[page_num] = {"page": page_num + 1, "error": str(e)}

    # Ghép kết quả theo đúng thứ tự từ đầu đến cuối
    full_final_result = ""
    for i in range(total_pages):
        if results_map[i].strip():
            full_final_result += results_map[i] + "\n\n---\n\n"

    vision_debug = [debug_map.get(i, {"page": i + 1}) for i in range(total_pages)]
    return full_final_result, vision_debug


def refine_text_with_ai(current_text, refinement_prompt, image_bytes=None):
    model = build_gemini_model(temperature=0.1)
    if image_bytes:
        img_b64 = base64.b64encode(image_bytes).decode("utf-8")
        content = [
            {
                "type": "text",
                "text": f"{EDUCATIONAL_SYSTEM_PROMPT}\n\nBẢN THẢO HIỆN TẠI:\n<draft>\n{current_text}\n</draft>\n\nYÊU CẦU SỬA LỖI:\n{refinement_prompt}\n\nHãy nhìn ảnh đính kèm để căn chỉnh lại cho giống với khoảng cách/định dạng trong ảnh.",
            },
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{img_b64}"},
            },
        ]
        return clean_ai_output(model.invoke([HumanMessage(content=content)]).content)
    else:
        full_prompt = f"{EDUCATIONAL_SYSTEM_PROMPT}\n\nBẢN THẢO HIỆN TẠI:\n<draft>\n{current_text}\n</draft>\n\nYÊU CẦU SỬA LỖI:\n{refinement_prompt}"
        return clean_ai_output(model.invoke(full_prompt).content)


def is_service_unavailable_error(error):
    """Detect Google 503-like errors without requiring google.api_core at import time."""
    error_name = error.__class__.__name__
    message = str(error).lower()
    return error_name == "ServiceUnavailable" or (
        "503" in message and ("unavailable" in message or "overload" in message)
    )


# ==========================================
# 5. HÀM GHI FILE WORD (THUẬT TOÁN BẢNG TÀNG HÌNH CHIA ĐỀU)
# ==========================================
def preprocess_text_for_word(text):
    """Làm sạch text trước khi tạo Word: thu gọn gạch dưới dài, xóa dòng gạch dư."""
    if not text:
        return text
    # Thu gọn mọi chuỗi gạch dưới từ 4+ dấu thành 5 dấu (tránh giãn quá dài trong Word)
    text = re.sub(r"_{4,}", "_____", text)
    # Thu gọn chuỗi gạch ngang từ 4+ dấu thành 3 dấu
    text = re.sub(r"-{4,}", "---", text)
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        # Bỏ qua dòng chỉ chứa ký tự gạch/phân cách
        if (
            re.match(r"^[\s\-_.*=!]+$", stripped)
            and len(re.findall(r"[\-_.*=!]", stripped)) >= 2
        ):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def create_word_docx(processed_text):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    title = doc.add_heading("TÀI LIỆU ĐÃ ĐƯỢC AI BIÊN TẬP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Times New Roman"
    title.runs[0].font.color.rgb = RGBColor(30, 136, 229)
    doc.add_paragraph()

    processed_text = preprocess_text_for_word(processed_text)
    lines = processed_text.split("\n")
    current_table = None
    is_first_row = False

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            current_table = None
            continue

        # Bỏ qua các dòng chỉ chứa ký tự gạch/phân cách
        # (---, ___, ..., -----, _____, ........, _ _ _, - - -, v.v.)
        if (
            re.match(r"^[\s\-_.*=!]+$", line_stripped)
            and len(re.findall(r"[\-_.*=!]", line_stripped)) >= 2
        ):
            current_table = None
            continue

        # --- XỬ LÝ BẢNG TỪ VỰNG CÓ VIỀN ---
        if line_stripped.startswith("|") and line_stripped.endswith("|"):
            if re.match(r"^\|[\-\|\s:]+\|$", line_stripped):
                continue

            cols = [c.strip() for c in line_stripped.split("|")[1:-1]]

            if current_table is None:
                current_table = doc.add_table(rows=0, cols=len(cols))
                current_table.style = "Table Grid"
                is_first_row = True
            else:
                is_first_row = False

            row_cells = current_table.add_row().cells
            for i, col_text in enumerate(cols):
                if i < len(row_cells):
                    cell = row_cells[i]
                    p = cell.paragraphs[0]
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if is_first_row
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )

                    if is_first_row:
                        shading_elm = parse_xml(
                            r'<w:shd {} w:fill="1E88E5"/>'.format(nsdecls("w"))
                        )
                        cell._tc.get_or_add_tcPr().append(shading_elm)

                    parts = re.split(r"\*\*(.*?)\*\*", col_text)
                    for j, part in enumerate(parts):
                        run = p.add_run(part)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

                        if is_first_row:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.bold = True
                        else:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            if j % 2 != 0:
                                run.bold = True

        # --- TẠO BẢNG TÀNG HÌNH ĐỂ ÉP CỘT TRẮC NGHIỆM A, B, C, D ĐỀU NHAU ---
        elif (
            "**A.**" in line
            and "**B.**" in line
            and "**C.**" in line
            and "**D.**" in line
        ):
            current_table = None

            # 1. Tìm vị trí của các đáp án
            idx_a = line.find("**A.**")
            idx_b = line.find("**B.**")
            idx_c = line.find("**C.**")
            idx_d = line.find("**D.**")

            # 2. Xử lý phần "Question" nếu AI lỡ viết dính liền trên cùng 1 dòng
            question_part = line[:idx_a].strip()
            if question_part:
                p_question = doc.add_paragraph()
                q_parts = re.split(r"\*\*(.*?)\*\*", question_part)
                for j, q_part in enumerate(q_parts):
                    run = p_question.add_run(q_part)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    if j % 2 != 0:
                        run.bold = True

            # 3. Chia 4 đáp án
            parts_text = [
                line[idx_a:idx_b].strip(),
                line[idx_b:idx_c].strip(),
                line[idx_c:idx_d].strip(),
                line[idx_d:].strip(),
            ]

            # 4. Tạo bảng 1 hàng 4 cột (CHIA ĐỀU KÍCH THƯỚC)
            mc_table = doc.add_table(rows=1, cols=4)
            mc_table.autofit = False

            # Ép 4 cột rộng bằng nhau tăm tắp (1.6 inches mỗi cột)
            widths = [Inches(1.6), Inches(1.6), Inches(1.6), Inches(1.6)]
            for cell, width in zip(mc_table.rows[0].cells, widths):
                cell.width = width

            for i, p_text in enumerate(parts_text):
                cell = mc_table.cell(0, i)
                p = cell.paragraphs[0]

                bold_parts = re.split(r"\*\*(.*?)\*\*", p_text)
                for j, b_part in enumerate(bold_parts):
                    run = p.add_run(b_part)
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    if j % 2 != 0:
                        run.bold = True

        # --- XỬ LÝ VĂN BẢN BÌNH THƯỜNG ---
        else:
            current_table = None
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            parts = re.split(r"\*\*(.*?)\*\*", line)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.text = part
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if i % 2 != 0:
                    run.bold = True

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ==========================================
# 6. HÀM AI: TỰ ĐỘNG GIẢI ĐỀ & TRÍCH XUẤT JSON
# ==========================================
def extract_quiz_to_json(text):
    # Dùng 1.5-flash để tốc độ nhanh và chịu tải text dài tốt hơn
    model = build_gemini_model(temperature=0.1)

    prompt = f"""
    Đọc tài liệu sau và tìm TẤT CẢ các câu hỏi trắc nghiệm. 
    Nhiệm vụ của bạn là trích xuất dữ liệu theo CÁC QUY TẮC TUYỆT ĐỐI SAU:

    1. CÂU HỎI: 
       - NẾU LÀ CÂU HỘI THOẠI / ĐỌC HIỂU: Hãy tóm tắt lại hoặc gộp chúng thành 1 CÂU DUY NHẤT. 
       - NẾU LÀ BÀI ĐIỀN TỪ: Chỉ trích xuất đúng 1 câu văn chứa chỗ trống.
       - TUYỆT ĐỐI KHÔNG SỬ DỤNG KÝ TỰ XUỐNG DÒNG. Toàn bộ câu hỏi phải nằm trên 1 dòng thẳng tắp.
    2. ĐÁP ÁN: Trích xuất đúng 4 đáp án (XÓA BỎ các ký tự A., B., C., D. ở đầu).
    3. ĐÓNG VAI GIÁO VIÊN: Tự động giải đề để tìm ra đáp án đúng.
    
    TRẢ VỀ ĐÚNG ĐỊNH DẠNG JSON SAU (Không chứa text nào khác):
    [
        {{
            "question": "Nội dung câu hỏi (Tuyệt đối không có dấu xuống dòng)?",
            "answers": ["Đáp án 1", "Đáp án 2", "Đáp án 3", "Đáp án 4"],
            "correct_index": 1, 
            "correct_text": "COPY CHÍNH XÁC 100% text từ 1 trong 4 đáp án trên"
        }}
    ]
    
    TÀI LIỆU GỐC:
    {text}
    """

    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = clean_ai_output(model.invoke(prompt).content)
            clean_json = response.replace("```json", "").replace("```", "").strip()

            # Cố gắng dịch cục text thành JSON
            return json.loads(clean_json)

        except Exception as e:
            error_msg = str(e)

            # Nếu gặp lỗi Google quá tải, đi ngủ 15s rồi thử lại
            if "503" in error_msg or "429" in error_msg:
                if attempt < max_retries - 1:
                    time.sleep(15)
                    continue

            # NẾU LỖI LÀ DO AI VIẾT SAI FORMAT JSON HOẶC LỖI KHÁC BẤT NGỜ
            import streamlit as st

            st.error(f"❌ Lỗi trích xuất dữ liệu ở lần thử {attempt + 1}: {error_msg}")

            # Phơi bày nguyên văn những gì AI đã viết để bạn soi xem nó ngo nguậy chỗ nào
            if "response" in locals():
                with st.expander(
                    "👀 Bấm vào đây để xem AI đã viết rác gì khiến hệ thống lỗi"
                ):
                    st.text(response)

            return None

    return None


# ==========================================
# 7. HÀM PYTHON: TẠO FILE CHO KAHOOT & BLOOKET (SIÊU TỐC)
# ==========================================
def generate_edtech_files(quiz_json):
    kahoot_data = []

    blooket_io = io.StringIO()
    blooket_io.write("Blooket Dummy Title,,,,,,,\n")
    fieldnames = [
        "Question #",
        "Question Text",
        "Answer 1",
        "Answer 2",
        "Answer 3",
        "Answer 4",
        "Time Limit (sec)",
        "Correct Answer(s)",
    ]
    writer = csv.DictWriter(blooket_io, fieldnames=fieldnames, lineterminator="\n")
    writer.writeheader()

    # CHỈ DÙNG 1 VÒNG LẶP DUY NHẤT CHO CẢ 2 NỀN TẢNG
    for i, q in enumerate(quiz_json):
        ans_idx = q["correct_index"]
        ans_idx = 1 if ans_idx == 0 else (4 if ans_idx > 4 else ans_idx)

        # Dọn rác văn bản (Chỉ làm 1 lần)
        clean_q = str(q["question"]).replace("\n", " ").replace("\r", " ").strip()
        clean_a1 = str(q["answers"][0]).replace("\n", " ").strip()
        clean_a2 = str(q["answers"][1]).replace("\n", " ").strip()
        clean_a3 = str(q["answers"][2]).replace("\n", " ").strip()
        clean_a4 = str(q["answers"][3]).replace("\n", " ").strip()

        # 1. Đút dữ liệu vào mảng Kahoot
        kahoot_data.append(
            {
                "Question - max 120 characters": clean_q[:120],
                "Answer 1 - max 75 characters": clean_a1[:75],
                "Answer 2 - max 75 characters": clean_a2[:75],
                "Answer 3 - max 75 characters": clean_a3[:75],
                "Answer 4 - max 75 characters": clean_a4[:75],
                "Time limit (sec) - 5, 10, 20, 30, 60, 90, 120, or 240": 20,
                "Correct answer(s) - 1, 2, 3, or 4": ans_idx,
            }
        )

        # 2. Viết trực tiếp dữ liệu vào Blooket CSV
        writer.writerow(
            {
                "Question #": i + 1,
                "Question Text": clean_q,
                "Answer 1": clean_a1,
                "Answer 2": clean_a2,
                "Answer 3": clean_a3,
                "Answer 4": clean_a4,
                "Time Limit (sec)": 20,
                "Correct Answer(s)": ans_idx,
            }
        )

    # Đóng gói Excel cho Kahoot
    df_kahoot = pd.DataFrame(kahoot_data)
    kahoot_io = io.BytesIO()
    df_kahoot.to_excel(kahoot_io, index=False, engine="openpyxl")
    kahoot_io.seek(0)

    return kahoot_io, blooket_io.getvalue().encode("utf-8")


# ==========================================
# GIAO DIỆN STREAMLIT
# ==========================================
st.markdown(
    """
    <style>
    div[data-testid="stVerticalBlock"] > div { border-radius: 12px; }
    div[data-testid="stAlert"] { border-radius: 8px; font-weight: 500; }
    table { border-collapse: collapse; width: 100%; border-radius: 8px; overflow: hidden; }
    th { background-color: #1E88E5; color: white; }
    </style>
    """,
    unsafe_allow_html=True,
)

if "draft_text" not in st.session_state:
    st.session_state.draft_text = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

with st.sidebar:
    st.header("⚙️ Thiết lập ban đầu")
    uploaded_pdf = st.file_uploader("1. Tải tài liệu lên", type=["pdf"])

    # YÊU CẦU MẶC ĐỊNH (Người dùng có thể xóa đi gõ lại trên Web)
    default_user_prompt = """Biên tập tài liệu này theo đúng chuẩn form đề thi:
1. BẢNG TỪ VỰNG & BẢNG CẤU TRÚC: Giữ nguyên đầy đủ các cột, định dạng bằng markdown table.
2. QUIZ & GRAMMAR: Giữ nguyên nội dung, định dạng rõ ràng.
3. CÂU HỎI TRẮC NGHIỆM: Bôi đậm chữ Question, bôi đậm A, B, C, D và dàn đều trên 1 dòng nếu đáp án trắc nghiệm ngắn có thể dàn đủ trên 1 dòng, nếu đáp án trắc nghiệm dài hãy để mỗi đáp án trắc nghiệm 1 dòng.
4. ĐỌC HIỂU: Giữ nguyên đoạn văn liền mạch, chỗ điền từ dùng _____, giữ nguyên số câu ở trước (ví dụ (12)_______).
5. Giữ nguyên layout, format, font, nhất là các từ đang được in đậm hay gạch chân."""

    user_prompt = st.text_area(
        "2. Yêu cầu xử lý:", value=default_user_prompt, height=250
    )
    process_btn = st.button(
        "🚀 Bắt đầu tạo Bản Nháp", type="primary", use_container_width=True
    )

if process_btn:
    if not uploaded_pdf:
        st.sidebar.error("Bạn chưa tải file PDF lên!")
    else:
        should_rerun = False
        with st.status("Đang phân tích tài liệu...", expanded=True) as status:
            try:  # BẮT ĐẦU MẶC ÁO GIÁP
                st.write("Đang quét nội dung...")
                raw_text = extract_text_from_upload(uploaded_pdf)

                mode_used = "text" if raw_text.strip() else "vision"
                result = ""
                vision_debug = None

                if not raw_text.strip():
                    st.write("Phát hiện PDF ảnh scan! Sử dụng OCR + AI...")
                    images = extract_images_from_upload(uploaded_pdf)

                    if not images:
                        status.update(
                            label="Không render được ảnh từ PDF",
                            state="error",
                            expanded=True,
                        )
                        st.error(
                            "❌ Không chuyển PDF sang ảnh được. "
                            "Bạn thử mở lại PDF để chắc chắn file không lỗi."
                        )
                        raise RuntimeError("PDF_TO_IMAGE_EMPTY")

                    # Sử dụng OCR Tesseract để trích xuất text từ tất cả trang
                    st.write("Đang trích xuất text từ tất cả trang bằng OCR...")
                    try:
                        ocr_full_text = ocr_images_with_tesseract(images)
                    except Exception as ocr_err:
                        st.error(f"❌ Lỗi OCR: {ocr_err}")
                        ocr_full_text = ""

                    if str(ocr_full_text or "").strip():
                        st.write("Đang định dạng text OCR theo chuẩn đề thi...")
                        result = process_text_with_ai(ocr_full_text, user_prompt)
                        mode_used = "ocr+ai"
                    else:
                        st.warning(
                            "⚠️ OCR không đọc được chữ, chuyển sang Vision API..."
                        )
                        # Fallback sang Vision API nếu OCR thất bại
                        result, vision_debug = process_vision_with_ai(
                            images, user_prompt
                        )
                        st.session_state.last_vision_debug = vision_debug
                else:
                    st.write("Đang định dạng theo chuẩn đề thi...")
                    result = process_text_with_ai(raw_text, user_prompt)

                result_text = str(result or "")

                # Nếu là file scan mà Vision trả rỗng -> thử OCR local (Tesseract)
                if (
                    not raw_text.strip()
                    and not result_text.strip()
                    and "images" in locals()
                ):
                    st.write("AI Vision trả rỗng. Đang thử OCR local (Tesseract)...")
                    try:
                        ocr_text = ocr_images_with_tesseract(images)
                    except Exception as ocr_err:
                        # Để nhánh bên dưới hiển thị thông báo + hướng dẫn
                        ocr_text = ""
                        st.session_state.last_ocr_error = str(ocr_err)
                    else:
                        st.session_state.last_ocr_error = None

                    if str(ocr_text or "").strip():
                        st.write("Đang định dạng lại từ text OCR...")
                        formatted_from_ocr = process_text_with_ai(ocr_text, user_prompt)
                        formatted_text = str(formatted_from_ocr or "")

                        if formatted_text.strip():
                            result_text = formatted_text
                            mode_used = "ocr+ai"
                        else:
                            st.warning(
                                "⚠️ AI vẫn trả rỗng sau OCR. Hệ thống sẽ hiển thị text OCR thô để bạn lấy nội dung."
                            )
                            result_text = ocr_text
                            mode_used = "ocr_raw"

                if not result_text.strip():
                    # Gemini đôi khi trả về chuỗi rỗng (bị chặn/recitation hoặc lỗi im lặng).
                    # Đừng giả vờ thành công; rơi về text thô nếu có.
                    if raw_text.strip():
                        st.warning(
                            "⚠️ AI trả về nội dung rỗng. Hệ thống sẽ tạm hiển thị TEXT THÔ trích xuất từ PDF để bạn vẫn lấy được nội dung."
                        )
                        result_text = raw_text
                        st.session_state.chat_history = [
                            {
                                "role": "assistant",
                                "content": "⚠️ AI trả về rỗng, mình đang hiển thị text thô trích xuất từ PDF để bạn kiểm tra.",
                            }
                        ]
                        status.update(
                            label="AI trả rỗng → đã rơi về text thô",
                            state="complete",
                            expanded=False,
                        )
                        should_rerun = True
                    else:
                        status.update(label="AI trả rỗng", state="error", expanded=True)
                        st.error(
                            "❌ Không trích xuất được chữ từ PDF (có thể là file scan) và AI Vision cũng trả về rỗng. "
                            "Bạn thử cắt nhỏ PDF (1–3 trang) hoặc dùng OCR local (Tesseract)."
                        )

                        if not is_tesseract_available():
                            with st.expander("🛠️ Cài Tesseract để bật OCR local"):
                                st.markdown(
                                    "- Windows (khuyến nghị): `winget install -e --id UB-Mannheim.TesseractOCR`\n"
                                    "- Sau khi cài, mở terminal mới và chạy: `tesseract --version`\n"
                                    "- Nếu vẫn không nhận: đặt biến môi trường `TESSERACT_CMD` trỏ tới `tesseract.exe` (thường ở `C:/Program Files/Tesseract-OCR/tesseract.exe`)."
                                )

                        with st.expander("🔎 Debug (thông tin tối thiểu)"):
                            st.write(
                                {
                                    "mode_used": mode_used,
                                    "raw_text_len": len(raw_text or ""),
                                    "images_len": len(images)
                                    if "images" in locals()
                                    else None,
                                    "vision_debug": st.session_state.get(
                                        "last_vision_debug"
                                    ),
                                    "ocr_available": is_tesseract_available(),
                                    "ocr_error": st.session_state.get("last_ocr_error"),
                                }
                            )
                else:
                    st.session_state.draft_text = result_text.strip()
                    st.session_state.chat_history = [
                        {
                            "role": "assistant",
                            "content": "Tài liệu đã được định dạng đầy đủ. Sếp kiểm tra lại nhé!",
                        }
                    ]
                    status.update(
                        label="Hoàn tất xử lý!", state="complete", expanded=False
                    )
                    should_rerun = True

                if should_rerun:
                    st.session_state.draft_text = result_text.strip()

            except Exception as e:
                if str(e) == "PDF_TO_IMAGE_EMPTY":
                    pass
                elif is_service_unavailable_error(e):
                    status.update(label="Lỗi máy chủ AI", state="error", expanded=True)
                    st.error(
                        "🤖 Máy chủ Google Gemini hiện đang quá tải (Lỗi 503) hoặc file của bạn quá nặng. Vui lòng cắt nhỏ file PDF ra hoặc chờ 1 phút rồi thử lại nhé!"
                    )
                else:
                    status.update(label="Lỗi hệ thống", state="error", expanded=True)
                    st.error(f"❌ Có lỗi xảy ra trong quá trình xử lý: {e}")

        if should_rerun:
            st.rerun()

if st.session_state.draft_text:
    col_preview, col_chat = st.columns([1.2, 1], gap="large")

    with col_preview:
        st.subheader("📑 Bản thảo hiện tại")
        docx_file = create_word_docx(st.session_state.draft_text)
        st.download_button(
            label="📥 TẢI FILE WORD (.DOCX)",
            data=docx_file,
            file_name="DeThi_HoanThien.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # --- CẬP NHẬT GIAO DIỆN XUẤT FILE GAME ---
        st.divider()
        st.markdown("### 🎮 Xuất file Game (Kahoot/Blooket)")

        # Khởi tạo bộ nhớ tạm để giữ file không bị mất khi load lại trang
        if "game_files" not in st.session_state:
            st.session_state.game_files = None

        if st.button("🎲 Tự động Giải đề & Trích xuất File Game", type="secondary"):
            with st.spinner(
                "🤖 AI đang làm bài để tìm đáp án đúng và phân loại dữ liệu..."
            ):
                quiz_data = extract_quiz_to_json(st.session_state.draft_text)

                if quiz_data:
                    # Tạo file cực nhanh với hàm đã tối ưu
                    kahoot_file, blooket_file = generate_edtech_files(quiz_data)
                    # LƯU VÀO BỘ NHỚ
                    st.session_state.game_files = (kahoot_file, blooket_file)
                    st.success(
                        "Đã trích xuất thành công! Bạn có thể tải file bất cứ lúc nào bên dưới:"
                    )
                else:
                    st.error(
                        "❌ Không tìm thấy câu hỏi trắc nghiệm hoặc có lỗi xảy ra."
                    )

        # Hiển thị nút tải xuống MÀ KHÔNG BỊ PHỤ THUỘC VÀO NÚT BẤM BÊN TRÊN
        if st.session_state.game_files:
            k_file, b_file = st.session_state.game_files

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="🟣 Tải file Excel cho KAHOOT!",
                    data=k_file,
                    file_name="Kahoot_Template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            with col2:
                st.download_button(
                    label="🟦 Tải file CSV cho BLOOKET",
                    data=b_file,
                    file_name="Blooket_Template.csv",
                    mime="text/csv",
                    use_container_width=True,
                )

        # Làm sạch text trước khi hiển thị preview
        preview_text = preprocess_text_for_word(st.session_state.draft_text)
        st.container(height=500, border=True).markdown(preview_text)

    with col_chat:
        st.subheader("💬 Chat với Biên Tập Viên")
        chat_container = st.container(height=350, border=True)
        with chat_container:
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
                    if "image" in message and message["image"]:
                        st.image(message["image"], width=200)

        uploaded_chat_img = st.file_uploader(
            "📎 Kéo thả / Chọn ảnh minh họa lỗi vào đây", type=["png", "jpg", "jpeg"]
        )

        if prompt := st.chat_input("VD: Giãn cách đáp án A B C D ra như trong ảnh"):
            img_bytes = uploaded_chat_img.getvalue() if uploaded_chat_img else None
            user_msg_display = (
                prompt
                if not img_bytes
                else f"📎 *[Đã đính kèm ảnh minh họa]*\n\n{prompt}"
            )

            st.session_state.chat_history.append(
                {"role": "user", "content": user_msg_display, "image": img_bytes}
            )

            with chat_container:
                with st.chat_message("user"):
                    st.markdown(prompt)
                    if img_bytes:
                        st.image(img_bytes, width=200)

            with chat_container:
                with st.chat_message("assistant"):
                    with st.spinner("👀 Đang xử lý..."):
                        new_draft = refine_text_with_ai(
                            st.session_state.draft_text, prompt, img_bytes
                        )
                        st.session_state.draft_text = new_draft
                        reply = "Đã chỉnh sửa xong theo yêu cầu!"
                        st.markdown(reply)
                        st.session_state.chat_history.append(
                            {"role": "assistant", "content": reply}
                        )

            st.rerun()
elif not uploaded_pdf:
    st.info("👈 Hãy tải file PDF Đề thi lên thanh bên trái để bắt đầu!")
